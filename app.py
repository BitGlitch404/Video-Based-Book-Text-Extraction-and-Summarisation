import os, cv2, json, uuid, threading, numpy as np, pytesseract
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from fpdf import FPDF
from docx import Document
from werkzeug.utils import secure_filename
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import nltk

try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt')
    nltk.download('punkt_tab')

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
EXTRACTED_FOLDER = 'Extracted'
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

for f in [UPLOAD_FOLDER, EXTRACTED_FOLDER]: os.makedirs(f, exist_ok=True)

# Store job states in memory
jobs = {}

def generate_fast_summary(text_list, sentences_per_page=2, final_sentences=6):
    """
    text_list: List of strings (one per image/page)
    """
    summarizer = LsaSummarizer()
    page_snippets = []

    # 1. Summarize each page (2-3 lines)
    for page_text in text_list:
        if len(page_text.strip()) < 50: continue # Skip empty/short pages
        
        parser = PlaintextParser.from_string(page_text, Tokenizer("english"))
        summary = summarizer(parser.document, sentences_per_page)
        page_snippets.append(" ".join([str(s) for s in summary]))

    # 2. Cascade: Summarize the collection of snippets for the "Whole Book"
    combined_context = " ".join(page_snippets)
    if not combined_context:
        return "No sufficient text found to summarize."

    parser_final = PlaintextParser.from_string(combined_context, Tokenizer("english"))
    final_summary_obj = summarizer(parser_final.document, final_sentences)
    
    return " ".join([str(s) for s in final_summary_obj])


def process_ocr_task(job_id, image_paths, timestamps):
    try:
        pdf = FPDF()
        doc = Document()
        results = []
        raw_texts_for_ai = []
        total = len(image_paths)
        
        USABLE_WIDTH = 190

        for idx, img_path in enumerate(image_paths):
            img = cv2.imread(img_path)
            # --- Preprocessing ---
            img = cv2.resize(img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            processed = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 20)
            
            text = pytesseract.image_to_string(processed, config='--oem 1 --psm 3')
            
            # CLEANING STEP: Replace common problematic characters
            clean_text = text.strip()
            clean_text = clean_text.replace('\u2018', "'").replace('\u2019', "'") # Smart quotes
            clean_text = clean_text.replace('\u201c', '"').replace('\u201d', '"') # Smart double quotes
            clean_text = clean_text.replace('\u2013', '-').replace('\u2014', '-') # Dashes

            # FINAL SAFETY: Force to Latin-1, ignoring anything else
            safe_text = clean_text.encode('latin-1', 'ignore').decode('latin-1')

            raw_texts_for_ai.append(safe_text)

            # Get data as a dictionary for easier iteration
            data = pytesseract.image_to_data(processed, output_type=pytesseract.Output.DICT)
            
            # Group text by block_num (this preserves paragraph structure)
            blocks = {}
            for i in range(len(data['text'])):
                if int(data['conf'][i]) > 0: # Filter out noise
                    b_num = data['block_num'][i]
                    if b_num not in blocks:
                        blocks[b_num] = {'text': [], 'top': data['top'][i], 'height': data['height'][i]}
                    blocks[b_num]['text'].append(data['text'][i])

            # Clean up blocks: Join words into strings and calculate word counts
            final_blocks = []
            total_frame_words = 0
            for b in blocks.values():
                block_string = " ".join(b['text']).strip()
                if block_string:
                    word_count = len(block_string.split())
                    total_frame_words += word_count
                    final_blocks.append({
                        'text': block_string,
                        'top': b['top'],
                        'h': b['height'],
                        'word_count': word_count
                    })

            pdf.add_page()
            
            # --- Rule 1: Whole Frame is a Title ---
            if total_frame_words > 0 and total_frame_words < 5:
                full_text = " ".join([b['text'] for b in final_blocks])
                safe_text = full_text.encode('latin-1', 'ignore').decode('latin-1')
                
                # PDF: Massive and Centered
                pdf.set_font("Arial", 'B', 40)
                pdf.ln(60) # Move to middle of page
                pdf.multi_cell(0, 20, txt=safe_text, align='C')
                
                # Word
                h = doc.add_heading(safe_text, level=0)
                h.alignment = 1
            else:
                # --- Rule 2 & 3: Headings and Paragraphs ---
                for i, block in enumerate(final_blocks):
                    safe_block = block['text'].encode('latin-1', 'ignore').decode('latin-1')
                    
                    # Determine if this block is a heading based on gap
                    is_heading = False
                    if block['word_count'] < 10:
                        # Check vertical distance to previous/next block
                        prev_gap = (block['top'] - final_blocks[i-1]['top']) if i > 0 else 0
                        next_gap = (final_blocks[i+1]['top'] - block['top']) if i < len(final_blocks)-1 else 0
                        
                        threshold = block['h'] * 2.5 # Significant gap
                        if prev_gap > threshold or next_gap > threshold or i == 0:
                            is_heading = True

                    if is_heading:
                        pdf.set_font("Arial", 'B', 20) # Bold heading
                        pdf.ln(10)
                        pdf.multi_cell(0, 10, txt=safe_block, align='L')
                        doc.add_heading(safe_block, level=1)
                    else:
                        # Normal Paragraph
                        # Dynamic font size to fill page width
                        char_count = len(safe_block)
                        # Avoid division by zero and set reasonable scale
                        if char_count > 0:
                            calc_size = min(14, max(9, (USABLE_WIDTH / (char_count/2)) * 3))
                        else:
                            calc_size = 12
                            
                        pdf.set_font("Arial", size=calc_size)
                        pdf.multi_cell(0, 6, txt=safe_block, align='L')
                        pdf.ln(4) # Paragraph spacing
                        doc.add_paragraph(safe_block)

            if idx < total - 1:
                doc.add_page_break()
            
            jobs[job_id]["progress"] = int(((idx + 1) / total) * 100)
            results.append({"timestamp": f"{float(timestamps[idx]):.2f}", "text": text})

        pdf_output = os.path.join(UPLOAD_FOLDER, f"result_{job_id}.pdf")
        doc_output = os.path.join(UPLOAD_FOLDER, f"result_{job_id}.docx")

        pdf.output(pdf_output)
        doc.save(doc_output)

        final_summary = generate_fast_summary(raw_texts_for_ai)

        jobs[job_id].update({
            "status": "completed",
            "results": results,
            "summary": final_summary,
            "pdf_url": f"/api/download-pdf/{job_id}",
            "doc_url": f"/api/download-doc/{job_id}"
        })
    except Exception as e:
        jobs[job_id].update({"status": "error", "error": str(e)})
    
    print(f"--- Job {job_id} Finished. Results: {len(results)} ---")

@app.route('/api/extract-text', methods=['POST'])
def start_ocr():
    if 'images' not in request.files:
        return jsonify({"error": "No images uploaded"}), 400
    
    job_id = str(uuid.uuid4())
    images = request.files.getlist('images')
    timestamps = json.loads(request.form.get('timestamps', '[]'))
    
    saved_paths = []
    for img in images:
        path = os.path.join(EXTRACTED_FOLDER, f"{job_id}_{secure_filename(img.filename)}")
        img.save(path)
        saved_paths.append(path)

    jobs[job_id] = {"status": "processing", "progress": 0}
    
    # Run in background
    thread = threading.Thread(target=process_ocr_task, args=(job_id, saved_paths, timestamps))
    thread.start()

    return jsonify({"job_id": job_id})

@app.route('/api/job-status/<job_id>')
def get_status(job_id):
    status_data = jobs.get(job_id, {"status": "not_found"})
    print(f"[PY LOG] Job {job_id} requested. Current status: {status_data.get('status')} - Progress: {status_data.get('progress')}%")
    return jsonify(status_data)

@app.route('/api/download-pdf/<job_id>')
def download_pdf(job_id):
    return send_file(os.path.join(UPLOAD_FOLDER, f"result_{job_id}.pdf"), as_attachment=True)

@app.route('/api/download-doc/<job_id>')
def download_doc(job_id):
    return send_file(os.path.join(UPLOAD_FOLDER, f"result_{job_id}.docx"), as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=5000)